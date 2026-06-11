"""
Resume Converter v7.8 — Flask API
POST /convert  { "text": "...", "filename": "Name.docx" }
"""

from flask import Flask, request, send_file
from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import io, re

app = Flask(__name__)
FONT = "Calibri"

MULTI_WORD = [
    ('PROFESSIONAL EXPERIENCE', 'PROF_EXP_TAG'),
    ('WORK EXPERIENCE',         'WORK_EXP_TAG'),
    ('TECHNICAL SKILLS',        'TECH_SKILLS_TAG'),
    ('CORE COMPETENCIES',       'CORE_COMP_TAG'),
]
ALL_SECTIONS = ([real for real,_ in MULTI_WORD] +
    ['SUMMARY','SKILLS','PROJECTS','EDUCATION','CERTIFICATIONS',
     'AWARDS','PUBLICATIONS','VOLUNTEERING','LANGUAGES','INTERESTS',
     'OBJECTIVE','PROFILE','ACCOMPLISHMENTS','EXPERIENCE','REFERENCES'])
KNOWN_SECTIONS  = set(ALL_SECTIONS)
_SINGLE = ['PROF_EXP_TAG','WORK_EXP_TAG','TECH_SKILLS_TAG','CORE_COMP_TAG',
           'SUMMARY','SKILLS','PROJECTS','EDUCATION','CERTIFICATIONS',
           'AWARDS','PUBLICATIONS','VOLUNTEERING','LANGUAGES','INTERESTS',
           'OBJECTIVE','PROFILE','ACCOMPLISHMENTS','EXPERIENCE','REFERENCES']
_combined  = '|'.join(re.escape(s) for s in _SINGLE)
SECTION_RE = re.compile(r'(?<=[^\n\s])[ \t]+(' + _combined + r')(?=[ \t\n]|\Z)')

BULLET_SECTIONS = {'PROFESSIONAL EXPERIENCE','EXPERIENCE','WORK EXPERIENCE',
                   'PROJECTS','CERTIFICATIONS','ACCOMPLISHMENTS'}
SKILLS_SECTIONS = {'SKILLS','TECHNICAL SKILLS','CORE COMPETENCIES'}
PLAIN_SECTIONS  = {'SUMMARY','EDUCATION','LANGUAGES','INTERESTS',
                   'REFERENCES','OBJECTIVE','PROFILE'}

ACTION_VERBS = (
    'Accomplished','Accelerated','Achieved','Aligned','Analyzed','Applied',
    'Architected','Assessed','Audited','Automated',
    'Benchmarked','Built',
    'Classified','Clustered','Collaborated','Compared','Computed','Conducted',
    'Configured','Containerized','Contributed','Converted','Coordinated','Created',
    'Debugged','Defined','Delivered','Demonstrated','Deployed','Designed',
    'Detected','Developed','Documented',
    'Enabled','Enforced','Engineered','Enhanced','Ensured','Established',
    'Evaluated','Executed','Extracted',
    'Facilitated','Fine-tuned','Forecasted','Formulated',
    'Generated','Guided',
    'Identified','Implemented','Improved','Increased','Integrated','Introduced',
    'Launched','Led','Leveraged','Loaded',
    'Maintained','Managed','Measured','Mentored','Migrated','Minimized',
    'Modeled','Monitored',
    'Optimized','Orchestrated',
    'Packaged','Parsed','Performed','Piloted','Pioneered','Predicted',
    'Presented','Prioritized','Processed','Profiled','Prototyped','Provisioned',
    'Queried',
    'Reduced','Refactored','Released','Regularized','Researched','Resolved','Reviewed',
    'Scaled','Segmented','Simulated','Spearheaded','Standardized','Streamlined',
    'Supported',
    'Tested','Tracked','Trained','Transformed','Tuned',
    'Updated','Used','Utilized',
    'Validated','Verified','Visualized',
)

EMAIL_RE    = re.compile(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}')
LINKEDIN_RE = re.compile(r'(?:linkedin\.com/in/[\w-]+|Linkedin|LinkedIn)', re.I)
LINKEDIN_URL = 'https://www.linkedin.com/in/nikhil-reddy-a732b1404/'


def clean_text(text):
    text = re.sub(r'\*{1,3}([^*\n]+)\*{1,3}', r'\1', text)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\[[^\]]{0,60}\]', '', text)
    return text


def normalize(text):
    for real,tag in MULTI_WORD:
        w = real.split()
        text = re.sub(w[0]+r'[ \t]*\n[ \t]*'+w[1], real, text)
    for real,tag in MULTI_WORD: text = text.replace(real, tag)
    text = SECTION_RE.sub(r'\n\n\1', text)
    for real,tag in MULTI_WORD: text = text.replace(tag, real)
    for sec in ALL_SECTIONS:
        text = re.sub(r'^('+re.escape(sec)+r')[ \t]+([A-Za-z])',
                      r'\1\n\2', text, flags=re.MULTILINE)
    lines_out = []
    for line in text.split('\n'):
        if re.search(r'[A-Z][A-Z/|&\s]{10,}', line) and \
           re.search(r'@|gmail|yahoo|outlook|\+1', line):
            m = re.search(r'([A-Z][a-z]+(?: [A-Z][a-z]+)+)[ \t]+([A-Z][A-Z/|&\s]{10,})', line)
            if m:
                name = m.group(1).strip()
                rest = line[m.start(2):]
                cm   = re.search(r'([A-Z][a-z]+,[ \t]*[A-Z]{2}.*)', rest)
                if cm:
                    lines_out += [name, rest[:cm.start()].strip(), cm.group(1).strip()]
                    continue
        lines_out.append(line)
    text = '\n'.join(lines_out)
    return re.sub(r'\n{3,}', '\n\n', text).strip()


def preprocess(text):
    out = []
    for line in text.split('\n'):
        if re.match(r'^-\s{3}', line):
            out.append(('BULLET', line[4:].rstrip())); continue
        if re.match(r'^[-*]\s+', line):
            out.append(('BULLET', re.sub(r'^[-*]\s+','',line).strip())); continue
        if not line.strip():
            out.append(('BLANK', '')); continue
        out.append(('TEXT', line.strip()))
    return out


def is_section(text):
    s = text.strip().upper()
    if s in KNOWN_SECTIONS: return True
    if len(s) <= 35 and re.match(r'^[A-Z][A-Z\s&/]+$', s): return True
    return False

def is_job_line(text):
    s = text.strip()
    return ('|' in s and bool(
        re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}', s)
        or 'Present' in s))

def is_edu_degree(text):
    return any(k in text for k in (
        'Master','Bachelor','Doctor','PhD','BSc','MSc','MBA',
        'B.E','M.E','B.Tech','M.Tech','Science','Engineering'))

def looks_like_contact(text):
    s = text.strip()
    return bool('@' in s or '+1' in s or
        re.search(r'\b\d{3}[-.\s]\d{3}[-.\s]\d{4}\b', s) or
        re.search(r'linkedin|github|portfolio', s, re.I) or
        (('Tampa' in s or 'FL' in s or 'USA' in s) and '|' in s))


def set_font(run, size_pt, bold=False, color=None, underline=False):
    run.font.name    = FONT
    run.font.size    = Pt(size_pt)
    run.font.bold    = bold
    run.font.italic  = False
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    f   = rPr.find(qn('w:rFonts'))
    if f is None:
        f = OxmlElement('w:rFonts'); rPr.insert(0, f)
    for attr in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'):
        f.set(qn(attr), FONT)

def set_spacing(para, before=0, after=0, line=276):
    pPr = para._p.get_or_add_pPr()
    sp  = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing'); pPr.append(sp)
    sp.set(qn('w:before'),   str(before))
    sp.set(qn('w:after'),    str(after))
    sp.set(qn('w:line'),     str(line))
    sp.set(qn('w:lineRule'), 'auto')

def set_align(para, alignment):
    pPr = para._p.get_or_add_pPr()
    jc  = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc'); pPr.append(jc)
    jc.set(qn('w:val'), alignment)

def add_bottom_border(para):
    pPr = para._p.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),'auto')
    bdr.append(bot)
    sp = pPr.find(qn('w:spacing'))
    sp.addprevious(bdr) if sp is not None else pPr.append(bdr)

def keep_together(para, keep_lines=True, keep_next=False):
    pPr = para._p.get_or_add_pPr()
    if keep_lines:
        kl = pPr.find(qn('w:keepLines'))
        if kl is None:
            kl = OxmlElement('w:keepLines'); pPr.append(kl)
    if keep_next:
        kn = pPr.find(qn('w:keepNext'))
        if kn is None:
            kn = OxmlElement('w:keepNext'); pPr.append(kn)

def add_tab_stop(para, position_inches):
    pPr = para._p.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs'); pPr.append(tabs)
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(position_inches * 1440)))
    tabs.append(tab)

def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')
    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    run_elem.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run_elem.append(t)
    hyperlink.append(run_elem)
    para._p.append(hyperlink)
    return hyperlink


def make_para(doc, style='Normal', before=0, after=0, line=276, align='left', border=False):
    para = doc.add_paragraph(style=style)
    set_spacing(para, before, after, line)
    set_align(para, align)
    if border: add_bottom_border(para)
    keep_together(para, keep_lines=True, keep_next=False)
    return para

def add_run(para, text, size_pt, bold=False, color=None, underline=False):
    run = para.add_run(text)
    set_font(run, size_pt, bold, color, underline)
    return run


def p_spacer(doc):
    para = make_para(doc, style='Normal', before=0, after=0, line=200)
    run  = para.add_run(' ')
    set_font(run, 4, bold=False)
    return para


def p_name(doc, text):
    para = make_para(doc, before=0, after=0, line=220, align='center')
    add_run(para, text, 20, bold=True)
    return para

def p_subtitle(doc, text):
    para = make_para(doc, before=0, after=0, line=220, align='center')
    add_run(para, text, 12, bold=True)
    return para

def p_contact(doc, text):
    para = make_para(doc, before=0, after=0, line=220, align='center', border=True)
    parts = [p.strip() for p in text.split('|')]
    for idx, part in enumerate(parts):
        sep = ' | ' if idx > 0 else ''
        email_match    = EMAIL_RE.search(part)
        linkedin_match = LINKEDIN_RE.search(part)
        if email_match:
            if sep: add_run(para, sep, 10, bold=True)
            add_hyperlink(para, part, 'mailto:' + email_match.group())
        elif linkedin_match:
            if sep: add_run(para, sep, 10, bold=True)
            add_hyperlink(para, part, LINKEDIN_URL)
        else:
            add_run(para, sep + part, 10, bold=True)
    return para

def p_section(doc, text):
    para = make_para(doc, before=0, after=0, line=276, align='left', border=True)
    add_run(para, text, 11, bold=True)
    return para

def p_job(doc, text):
    para = make_para(doc, before=0, after=0, line=276, align='left')
    add_tab_stop(para, 7.77)
    date_match = re.search(
        r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}'
        r'\s*[–\-—]\s*(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}|Present))',
        text)
    if date_match:
        left_part = text[:date_match.start()].rstrip(' |–-—').strip()
        date_part = date_match.group(1).strip()
        add_run(para, left_part, 11, bold=True)
        add_run(para, '\t', 11, bold=False)
        add_run(para, date_part, 11, bold=True)
    else:
        add_run(para, text, 11, bold=True)
    return para

def p_bullet(doc, text):
    para = make_para(doc, style='List Bullet', before=0, after=0, line=276, align='left')
    add_run(para, text, 10, bold=False)
    return para

def p_skills(doc, text):
    para = make_para(doc, before=0, after=0, line=276, align='left')
    if ':' in text:
        label, rest = text.split(':', 1)
        add_run(para, label + ':', 10, bold=True)
        add_run(para, rest, 10, bold=False)
    else:
        add_run(para, text, 10, bold=False)
    return para

def p_summary(doc, text):
    para = make_para(doc, before=0, after=0, line=276, align='left')
    add_run(para, text, 10, bold=False)
    return para

def p_edu_degree(doc, text):
    para = make_para(doc, before=0, after=0, line=276, align='left')
    add_run(para, text, 10, bold=True)
    return para

def p_edu_school(doc, text):
    para = make_para(doc, before=0, after=0, line=276, align='left')
    add_run(para, text, 10, bold=False)
    return para

def p_plain(doc, text):
    para = make_para(doc, before=0, after=0, line=276, align='left')
    add_run(para, text, 10, bold=False)
    return para


def build_docx(text):
    text = clean_text(text)
    text = normalize(text)

    doc = Document()
    for para in doc.paragraphs:
        para._p.getparent().remove(para._p)

    sec = doc.sections[0]
    def twip(inches): return Emu(int(inches * 914400))
    sec.page_width    = twip(8.27)
    sec.page_height   = twip(11.69)
    sec.top_margin    = twip(0.25)
    sec.bottom_margin = twip(0.25)
    sec.left_margin   = twip(0.25)
    sec.right_margin  = twip(0.25)

    lines           = preprocess(text)
    stage           = 'name'
    current_section = None
    current_block   = []
    first_section   = True
    project_count   = 0

    def close_block():
        n = len(current_block)
        for i, p in enumerate(current_block):
            keep_together(p, keep_lines=True, keep_next=(i < n - 1))
        current_block.clear()

    for kind, raw in lines:
        stripped  = raw.strip()
        is_bullet = (kind == 'BULLET')

        if kind == 'BLANK' or not stripped:
            continue

        if stage == 'name':
            p_name(doc, stripped); stage = 'subtitle'; continue

        if stage == 'subtitle':
            if re.match(r'^[A-Z][A-Z/|&\s]+$', stripped) and '|' in stripped:
                p_subtitle(doc, stripped); stage = 'contact'; continue
            else:
                stage = 'contact'

        if stage == 'contact':
            if is_section(stripped):
                stage = 'body'
            elif stripped.lower() in ('linkedin', 'github', 'portfolio'):
                continue
            elif looks_like_contact(stripped):
                p_contact(doc, stripped); continue
            elif not is_section(stripped):
                p_contact(doc, stripped); continue

        stage = 'body'

        if is_section(stripped):
            close_block()
            if first_section:
                first_section = False
            else:
                p_spacer(doc)
            project_count   = 0
            current_section = stripped.upper()
            para = p_section(doc, stripped)
            keep_together(para, keep_lines=True, keep_next=True)
            current_block = []
            continue

        sec_name = current_section or ''

        if sec_name == 'SUMMARY':
            para = p_summary(doc, stripped)
            keep_together(para, keep_lines=True, keep_next=False)
            continue

        if sec_name in SKILLS_SECTIONS:
            para = p_skills(doc, stripped)
            keep_together(para, keep_lines=True, keep_next=False)
            continue

        if sec_name == 'EDUCATION':
            para = (p_edu_degree if is_edu_degree(stripped) else p_edu_school)(doc, stripped)
            current_block.append(para)
            continue

        if sec_name in PLAIN_SECTIONS - {'SUMMARY', 'EDUCATION'}:
            para = p_plain(doc, stripped)
            keep_together(para, keep_lines=True, keep_next=False)
            continue

        if sec_name in ('PROFESSIONAL EXPERIENCE', 'EXPERIENCE', 'WORK EXPERIENCE'):
            if is_job_line(stripped):
                close_block()
                para = p_job(doc, stripped)
                current_block = [para]
            else:
                para = p_bullet(doc, stripped)
                current_block.append(para)
            continue

        if sec_name == 'PROJECTS':
            if is_bullet or stripped.startswith(ACTION_VERBS):
                para = p_bullet(doc, stripped)
                current_block.append(para)
            else:
                close_block()
                if project_count > 0:
                    p_spacer(doc)
                project_count += 1
                para = p_job(doc, stripped)
                current_block = [para]
            continue

        if sec_name in BULLET_SECTIONS:
            para = p_bullet(doc, stripped)
            current_block.append(para)
            continue

        para = p_plain(doc, stripped)
        keep_together(para, keep_lines=True, keep_next=False)

    close_block()
    return doc


# Flask routes

@app.route('/convert', methods=['POST'])
def convert():
    data     = request.json or {}
    text     = data.get('text', '')
    filename = data.get('filename', 'resume.docx')
    if not filename.endswith('.docx'): filename += '.docx'
    doc = build_docx(text)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return send_file(buf, as_attachment=True, download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/health', methods=['GET'])
def health():
    return {'status': 'ok', 'version': '7.8'}

@app.route('/text-to-pdf', methods=['POST'])
def text_to_pdf():
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.utils import simpleSplit

    optimized_text = request.form.get('optimized_text', '')
    if not optimized_text:
        return {'error': 'No optimized_text provided'}, 400

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    margin = 50
    x = margin
    y = height - margin
    line_height = 14
    max_width = width - 2 * margin
    c.setFont("Helvetica", 11)

    for line in optimized_text.split('\n'):
        line = line.strip()
        if not line:
            y -= line_height / 2
            continue
        wrapped = simpleSplit(line, "Helvetica", 11, max_width)
        for wline in wrapped:
            if y < margin:
                c.showPage()
                c.setFont("Helvetica", 11)
                y = height - margin
            c.drawString(x, y, wline)
            y -= line_height

    c.save()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
        download_name='optimized_resume.pdf',
        mimetype='application/pdf')

@app.route('/extract', methods=['POST'])
def extract():
    if 'file' not in request.files:
        return {'error': 'No file'}, 400
    file = request.files['file']
    filename = file.filename.lower()
    content = file.read()
    if filename.endswith('.pdf') or content[:4] == b'%PDF':
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            text = '\n'.join([page.extract_text() or '' for page in pdf.pages])
    else:
        doc = Document(io.BytesIO(content))
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    return {'text': text, 'filename': file.filename}
